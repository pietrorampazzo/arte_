
import os
import sys
import pandas as pd
import zipfile
import tempfile
import shutil
import re
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any
import time
import traceback

# Importações para extração de documentos
try:
    import fitz  # PyMuPDF
    import camelot
    from pdf2image import convert_from_path
    import pytesseract
    CAMELOT_AVAILABLE = True
    PDF_TOOLS_AVAILABLE = True
except ImportError:
    print("⚠️ AVISO: Ferramentas de PDF (PyMuPDF, camelot, pdf2image, pytesseract) não estão completamente instaladas. A extração de PDF será limitada.")
    CAMELOT_AVAILABLE = False
    PDF_TOOLS_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    print("⚠️ AVISO: python-docx não está instalado. Arquivos DOCX não serão processados.")
    DOCX_AVAILABLE = False

try:
    import chardet
    CHARDET_AVAILABLE = True
except ImportError:
    print("⚠️ AVISO: chardet não está instalado. Detecção automática de encoding será limitada.")
    CHARDET_AVAILABLE = False

# ======================================================================
# CONFIGURAÇÕES E CONSTANTES
# ======================================================================

# --- Caminhos dos Arquivos ---
BASE_DIR = r"C:\Users\pietr\OneDrive\.vscode\arte_"
EDITAIS_DIR = os.path.join(BASE_DIR, "EDITAIS")
OUTPUT_FILE = os.path.join(BASE_DIR, "DOWNLOADS", "consolidado.txt")

# --- Configurações de Ferramentas Externas ---
POPLER_PATH = os.path.join(BASE_DIR, "scripts", "Release-25.07.0-0", "poppler-25.07.0", "Library", "bin")
TESSERACT_CMD = os.path.join(BASE_DIR, "scripts", "Tesseract-OCR", "tesseract.exe")
if PDF_TOOLS_AVAILABLE:
    pytesseract.pytesseract.tesseract_cmd = str(TESSERACT_CMD)

# --- Extensões Suportadas ---
SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt', '.zip'}

# --- Opções de Encoding ---
ENCODING_OPTIONS = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

# ======================================================================
# FUNÇÕES AUXILIARES
# ======================================================================

def log_info(message: str) -> None:
    """Registra mensagem informativa."""
    print(f"[INFO] {message}")

def log_success(message: str) -> None:
    """Registra mensagem de sucesso."""
    print(f"[SUCCESS] {message}")

def log_warning(message: str) -> None:
    """Registra mensagem de aviso."""
    print(f"[WARNING] {message}")

def log_error(message: str) -> None:
    """Registra mensagem de erro."""
    print(f"[ERROR] {message}")

def log_debug(message: str) -> None:
    """Registra mensagem de debug."""
    print(f"[DEBUG] {message}")

# ======================================================================
# FUNÇÕES DE EXTRAÇÃO
# ======================================================================

def percorrer_diretorios(base_dir: str) -> List[Dict[str, str]]:
    """
    Percorre recursivamente todos os subdiretórios e retorna lista de arquivos
    com seus caminhos completos e tipos.
    """
    log_info(f"Iniciando varredura do diretório: {base_dir}")
    
    if not os.path.exists(base_dir):
        log_error(f"Diretório não encontrado: {base_dir}")
        return []
    
    arquivos_encontrados = []
    
    try:
        for root, dirs, files in os.walk(base_dir):
            for file in files:
                file_path = os.path.join(root, file)
                file_ext = Path(file).suffix.lower()
                
                if file_ext in SUPPORTED_EXTENSIONS:
                    arquivo_info = {
                        'caminho_completo': file_path,
                        'nome_arquivo': file,
                        'extensao': file_ext,
                        'diretorio': root,
                        'tamanho': os.path.getsize(file_path) if os.path.exists(file_path) else 0
                    }
                    arquivos_encontrados.append(arquivo_info)
        
        log_success(f"Encontrados {len(arquivos_encontrados)} arquivos suportados")
        
    except Exception as e:
        log_error(f"Erro ao percorrer diretórios: {str(e)}")
        return []
    
    return arquivos_encontrados

def extrair_texto_pdf(file_path: str) -> Dict[str, Any]:
    """
    Extrai texto de PDFs usando uma abordagem em duas etapas:
    1. Tenta extrair texto diretamente com PyMuPDF (fitz).
    2. Se o texto for insuficiente, usa OCR com Tesseract.
    Também extrai tabelas com camelot.
    """
    resultado = {'texto': '', 'tabelas': [], 'erro': None}
    
    if not PDF_TOOLS_AVAILABLE:
        resultado['erro'] = "Ferramentas de PDF não estão instaladas."
        return resultado

    texto_completo = ""

    try:
        log_info(f"Extraindo PDF: {os.path.basename(file_path)}")

        # Etapa 1: Extração de texto com PyMuPDF
        try:
            with fitz.open(file_path) as doc:
                for page in doc:
                    texto_completo += page.get_text()
            log_info("Texto extraído com PyMuPDF.")
        except Exception as e:
            log_warning(f"Falha na extração com PyMuPDF: {e}. Tentando OCR.")
            texto_completo = ""

        # Etapa 2: Se o texto for insuficiente, usar OCR
        if len(texto_completo.strip()) < 100: # Heurística: se tem menos de 100 caracteres, provavelmente é imagem
            log_info("Texto insuficiente, iniciando OCR com Tesseract...")
            try:
                paginas = convert_from_path(file_path, dpi=300, poppler_path=POPLER_PATH)
                texto_ocr = ""
                for i, pagina in enumerate(paginas):
                    log_debug(f"Lendo página {i+1}/{len(paginas)} (OCR)...")
                    texto_ocr += pytesseract.image_to_string(pagina, lang="por") + "\n\n"
                texto_completo = texto_ocr
                log_success("Texto extraído com OCR.")
            except Exception as e:
                log_error(f"Erro durante o OCR: {e}")
                # Não retorna aqui, pode ser que o texto do PyMuPDF (mesmo que pouco) seja útil

        resultado['texto'] = texto_completo

        # Manter a extração de tabelas com Camelot, que é uma boa ferramenta para isso
        if CAMELOT_AVAILABLE:
            try:
                log_info("Iniciando extração de tabelas com Camelot...")
                tables = camelot.read_pdf(file_path, flavor='lattice', edge_tol=500)
                if len(tables) == 0:
                    tables = camelot.read_pdf(file_path, flavor='stream')
                
                for table in tables:
                    if not table.df.empty:
                        resultado['tabelas'].append(table.df)
                log_success(f"Extraídas {len(resultado['tabelas'])} tabelas com Camelot.")
            except Exception as e:
                log_warning(f"Erro na extração de tabelas com camelot: {str(e)}")

    except Exception as e:
        resultado['erro'] = str(e)
        log_error(f"Erro ao extrair PDF {file_path}: {str(e)}")

    return resultado


def extrair_texto_docx(file_path: str) -> Dict[str, Any]:
    """
    Usa python-docx para extrair texto e tabelas de arquivos DOCX.
    """
    resultado = {'texto': '', 'tabelas': [], 'erro': None}
    
    if not DOCX_AVAILABLE:
        resultado['erro'] = "python-docx não está disponível"
        return resultado
    
    try:
        log_info(f"Extraindo DOCX: {os.path.basename(file_path)}")
        
        doc = Document(file_path)
        
        texto_completo = ""
        for paragraph in doc.paragraphs:
            texto_completo += paragraph.text + "\n"
        
        resultado['texto'] = texto_completo
        
        for table in doc.tables:
            dados_tabela = []
            for row in table.rows:
                linha_dados = [cell.text.strip() for cell in row.cells]
                dados_tabela.append(linha_dados)
            
            if dados_tabela:
                try:
                    df_tabela = pd.DataFrame(dados_tabela[1:], columns=dados_tabela[0])
                    resultado['tabelas'].append(df_tabela)
                except:
                    df_tabela = pd.DataFrame(dados_tabela)
                    resultado['tabelas'].append(df_tabela)
        
    except Exception as e:
        resultado['erro'] = str(e)
        log_error(f"Erro ao extrair DOCX {file_path}: {str(e)}")
    
    return resultado

def extrair_texto_doc(file_path: str) -> Dict[str, Any]:
    """
    Tenta lidar com arquivos DOC legados.
    """
    resultado = {'texto': '', 'tabelas': [], 'erro': None}
    log_warning(f"Arquivo DOC detectado: {os.path.basename(file_path)}. Formato não suportado diretamente.")
    resultado['erro'] = "Formato DOC não suportado"
    return resultado

def extrair_texto_excel(file_path: str) -> Dict[str, Any]:
    """
    Usa pandas para ler arquivos Excel.
    """
    resultado = {'texto': '', 'tabelas': [], 'erro': None}
    
    try:
        log_info(f"Extraindo Excel: {os.path.basename(file_path)}")
        
        excel_file = pd.ExcelFile(file_path)
        
        for sheet_name in excel_file.sheet_names:
            try:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                if not df.empty:
                    resultado['tabelas'].append(df)
                    resultado['texto'] += f"PLANILHA: {sheet_name}\n"
            except Exception as e:
                log_warning(f"Erro ao ler planilha '{sheet_name}': {str(e)}")
        
    except Exception as e:
        resultado['erro'] = str(e)
        log_error(f"Erro ao extrair Excel {file_path}: {str(e)}")
    
    return resultado

def extrair_texto_txt(file_path: str) -> Dict[str, Any]:
    """
    Leitura de arquivos de texto com detecção de encoding.
    """
    resultado = {'texto': '', 'tabelas': [], 'erro': None}
    
    try:
        log_info(f"Extraindo TXT: {os.path.basename(file_path)}")
        
        encoding_usado = 'utf-8'
        if CHARDET_AVAILABLE:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding_detectado = chardet.detect(raw_data)
                if encoding_detectado['encoding']:
                    encoding_usado = encoding_detectado['encoding']
        
        for encoding in [encoding_usado] + ENCODING_OPTIONS:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    resultado['texto'] = f.read()
                    break
            except UnicodeDecodeError:
                continue
        
        if not resultado['texto']:
            resultado['erro'] = "Não foi possível decodificar o arquivo"
        
    except Exception as e:
        resultado['erro'] = str(e)
        log_error(f"Erro ao extrair TXT {file_path}: {str(e)}")
    
    return resultado

def processar_arquivo_zip(file_path: str) -> Dict[str, Any]:
    """
    Extrai conteúdo de ZIP para diretório temporário e processa arquivos internos.
    """
    resultado = {'texto': '', 'tabelas': [], 'erro': None, 'arquivos_processados': []}
    
    try:
        log_info(f"Processando ZIP: {os.path.basename(file_path)}")
        
        with tempfile.TemporaryDirectory() as temp_dir:
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            arquivos_extraidos = percorrer_diretorios(temp_dir)
            
            for arquivo in arquivos_extraidos:
                caminho_arquivo = arquivo['caminho_completo']
                resultado_arquivo = extrair_conteudo_arquivo(caminho_arquivo)
                
                if resultado_arquivo and not resultado_arquivo.get('erro'):
                    resultado['texto'] += f"\n--- ARQUIVO ZIPADO: {arquivo['nome_arquivo']} ---\n"
                    resultado['texto'] += resultado_arquivo.get('texto', '')
                    resultado['tabelas'].extend(resultado_arquivo.get('tabelas', []))
                    resultado['arquivos_processados'].append(arquivo['nome_arquivo'])
        
    except Exception as e:
        resultado['erro'] = str(e)
        log_error(f"Erro ao processar ZIP {file_path}: {str(e)}")
    
    return resultado

def extrair_conteudo_arquivo(file_path: str) -> Dict[str, Any]:
    """
    Função que despacha para o extrator apropriado baseado na extensão do arquivo.
    """
    if not os.path.exists(file_path):
        return {'erro': f"Arquivo não encontrado: {file_path}"}
    
    extensao = Path(file_path).suffix.lower()
    
    if extensao == '.pdf':
        return extrair_texto_pdf(file_path)
    elif extensao == '.docx':
        return extrair_texto_docx(file_path)
    elif extensao == '.doc':
        return extrair_texto_doc(file_path)
    elif extensao in ['.xlsx', '.xls']:
        return extrair_texto_excel(file_path)
    elif extensao == '.txt':
        return extrair_texto_txt(file_path)
    elif extensao == '.zip':
        return processar_arquivo_zip(file_path)
    else:
        return {'erro': f"Extensão não suportada: {extensao}"}

# ======================================================================
# FUNÇÃO PRINCIPAL
# ======================================================================

def main():
    """
    Função principal que executa a extração e consolidação dos dados.
    """
    print("=" * 80)
    print(">> CONSOLIDADOR DE DADOS DE EDITAIS")
    print("=" * 80)
    
    if not os.path.exists(EDITAIS_DIR):
        log_error(f"Diretório de editais não encontrado: {EDITAIS_DIR}")
        return

    arquivos_a_processar = percorrer_diretorios(EDITAIS_DIR)
    
    if not arquivos_a_processar:
        log_warning("Nenhum arquivo encontrado para processar.")
        return

    conteudo_consolidado = []
    
    for arquivo_info in arquivos_a_processar:
        caminho = arquivo_info['caminho_completo']
        nome_arquivo = arquivo_info['nome_arquivo']
        
        log_info(f"Processando arquivo: {nome_arquivo}")
        
        resultado_extracao = extrair_conteudo_arquivo(caminho)
        
        if resultado_extracao.get('erro'):
            log_warning(f"Erro ao processar {nome_arquivo}: {resultado_extracao['erro']}")
            continue
            
        conteudo_consolidado.append(f"\n{'='*20} INÍCIO DO ARQUIVO: {nome_arquivo} {'='*20}\n")
        
        texto = resultado_extracao.get('texto', '')
        if texto:
            conteudo_consolidado.append(texto)
            
        tabelas = resultado_extracao.get('tabelas', [])
        if tabelas:
            conteudo_consolidado.append("\n--- TABELAS EXTRAÍDAS ---\n")
            for i, tabela in enumerate(tabelas):
                conteudo_consolidado.append(f"\n--- Tabela {i+1} ---\n")
                conteudo_consolidado.append(tabela.to_string())

        conteudo_consolidado.append(f"\n{'='*20} FIM DO ARQUIVO: {nome_arquivo} {'='*20}\n")

    try:
        with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
            f.write("\n".join(conteudo_consolidado))
        log_success(f"Arquivo consolidado salvo em: {OUTPUT_FILE}")
    except Exception as e:
        log_error(f"Erro ao salvar o arquivo consolidado: {str(e)}")

if __name__ == "__main__":
    main()

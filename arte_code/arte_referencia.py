#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
EXTRATOR DE EDITAIS COMPLETO
============================

Script principal para extrair descrições complementares de editais em diversos formatos
e gerar o arquivo master_heavy.xlsx com informações enriquecidas.

Autor: Traycer.AI
Data: SETEMBRO DE 2025
Versão: 1.0
"""

import os
import sys
import pandas as pd
import zipfile
import tempfile
import shutil
import re
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any
import time
import traceback

# Importações para extração de documentos
try:
    import pdfplumber
    import camelot
    CAMELOT_AVAILABLE = True
except ImportError:
    print("⚠️ AVISO: camelot-py não está instalado. Extração de tabelas PDF será limitada.")
    CAMELOT_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    print("⚠️ AVISO: python-docx não está instalado. Arquivos DOCX não serão processados.")
    DOCX_AVAILABLE = False

try:
    import chardet
    CHARDET_AVAILABLE = True
except ImportError:
    print("⚠️ AVISO: chardet não está instalado. Detecção automática de encoding será limitada.")
    CHARDET_AVAILABLE = False

# ======================================================================
# CONFIGURAÇÕES E CONSTANTES
# ======================================================================

# --- Caminhos dos Arquivos ---
BASE_DIR = r"C:\Users\pietr\OneDrive\.vscode\arte_"
EDITAIS_DIR = os.path.join(BASE_DIR, "DOWNLOADS", "EDITAIS_TESTE")
MASTER_XLSX_PATH = os.path.join(BASE_DIR, "EDITAIS", "master.xlsx")
OUTPUT_PATH = os.path.join(BASE_DIR, "DOWNLOADS", "ORCAMENTOS", "master_heavy.xlsx")

# --- Extensões Suportadas ---
SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt', '.zip'}

# --- Padrões de Regex para Detecção ---
COLUMN_PATTERNS = {
    'numero_item': [
        r'n[úu]mero?\s*(?:do\s*)?item', # "numero do item", "número item"
        r'item\s*n[úu]mero?', # "item numero"
        r'^n[°º]?$', # "nº", "n°"
        r'^item$', # "Item"
        r'seq[uü]?[eê]ncia', # "sequência"
        r'ordem', # "ordem"
        r'siafisico' # "Siafisico"
    ],
    'descricao': [
        r'descri[çc][ãa]o', # "descrição"
        r'especifica[çc][ãa]o', # "especificação"
        r'detalhamento', # "detalhamento"
        r'produto', # "produto", "nome do produto"
        r'material', # "material"
        r'objeto', # "objeto"
        r'catmat', # "CatMat"
        r'descri[çc][ãa]o\s*(?:do\s*)?(?:produto|material|item)' # "descrição do produto"
    ]
}

TABLE_KEYWORDS = [
    'condi[çc][õo]es gerais',
    'modelo de planilha',
    'especifica[çc][õo]es',
    'descri[çc][ãa]o detalhada',
    'memorial descritivo',
    'termo de refer[eê]ncia'
]

# --- Configurações do Camelot ---
CAMELOT_SETTINGS = {
    'flavor': 'lattice',
    'edge_tol': 500,
    'row_tol': 10,
    'column_tol': 0
}

# --- Opções de Encoding ---
ENCODING_OPTIONS = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

# ======================================================================
# FUNÇÕES AUXILIARES
# ======================================================================

def log_info(message: str, emoji: str = "ℹ️") -> None:
    """Registra mensagem informativa com emoji."""
    print(f"{emoji} {message}")

def log_success(message: str) -> None:
    """Registra mensagem de sucesso."""
    print(f"✅ {message}")

def log_warning(message: str) -> None:
    """Registra mensagem de aviso."""
    print(f"⚠️ {message}")

def log_error(message: str) -> None:
    """Registra mensagem de erro."""
    print(f"❌ {message}")

def log_debug(message: str) -> None:
    """Registra mensagem de debug."""
    print(f"🔍 DEBUG: {message}")

# ======================================================================
# FUNÇÕES PRINCIPAIS
# ======================================================================

def percorrer_diretorios(base_dir: str) -> List[Dict[str, str]]:
    """
    Percorre recursivamente todos os subdiretórios e retorna lista de arquivos
    com seus caminhos completos e tipos.
    
    Args:
        base_dir: Diretório base para busca
        
    Returns:
        Lista de dicionários com informações dos arquivos
    """
    log_info(f"Iniciando varredura do diretório: {base_dir}", "🔍")
    
    if not os.path.exists(base_dir):
        log_error(f"Diretório não encontrado: {base_dir}")
        return []
    
    arquivos_encontrados = []
    
    try:
        for root, dirs, files in os.walk(base_dir):
            for file in files:
                file_path = os.path.join(root, file)
                file_ext = Path(file).suffix.lower()
                
                if file_ext in SUPPORTED_EXTENSIONS:
                    arquivo_info = {
                        'caminho_completo': file_path,
                        'nome_arquivo': file,
                        'extensao': file_ext,
                        'diretorio': root,
                        'tamanho': os.path.getsize(file_path) if os.path.exists(file_path) else 0
                    }
                    arquivos_encontrados.append(arquivo_info)
        
        log_success(f"Encontrados {len(arquivos_encontrados)} arquivos suportados")
        
        # Estatísticas por tipo
        stats = {}
        for arquivo in arquivos_encontrados:
            ext = arquivo['extensao']
            stats[ext] = stats.get(ext, 0) + 1
        
        for ext, count in stats.items():
            log_info(f"  {ext}: {count} arquivo(s)")
            
    except Exception as e:
        log_error(f"Erro ao percorrer diretórios: {str(e)}")
        return []
    
    return arquivos_encontrados

def normalizar_numero_item(item_num: Any) -> str:
    """
    Padroniza números de itens removendo zeros à esquerda, espaços e
    lidando com diferentes formatos.
    
    Args:
        item_num: Número do item em qualquer formato
        
    Returns:
        Número normalizado como string
    """
    if pd.isna(item_num) or item_num is None:
        return ""
    
    # Converter para string
    item_str = str(item_num).strip()
    
    # Remover pontos finais e outros caracteres
    item_str = re.sub(r'[.\-_\s]+$', '', item_str)
    item_str = re.sub(r'^[.\-_\s]+', '', item_str)
    
    # Extrair apenas números
    numeros = re.findall(r'\d+', item_str)
    if numeros:
        # Pegar o primeiro número encontrado e remover zeros à esquerda
        return str(int(numeros[0]))
    
    return item_str

def detectar_colunas_tabela(df: pd.DataFrame) -> Dict[str, str]:
    """
    Usa heurísticas e regex para identificar colunas contendo números de itens
    e descrições em tabelas extraídas.
    
    Args:
        df: DataFrame da tabela extraída
        
    Returns:
        Dicionário com mapeamento de tipos para nomes de colunas
    """
    colunas_detectadas = {'numero': None, 'descricao': None}
    
    if df.empty:
        return colunas_detectadas
    
    # Analisar nomes das colunas
    for col in df.columns:
        col_lower = str(col).lower()
        
        # Detectar coluna de número
        for pattern in COLUMN_PATTERNS['numero_item']:
            if re.search(pattern, col_lower, re.IGNORECASE):
                colunas_detectadas['numero'] = col
                break
        
        # Detectar coluna de descrição
        for pattern in COLUMN_PATTERNS['descricao']:
            if re.search(pattern, col_lower, re.IGNORECASE):
                colunas_detectadas['descricao'] = col
                break
    
    # Se não encontrou pelos nomes, analisar conteúdo
    if not colunas_detectadas['numero']:
        for col in df.columns:
            # Verificar se a coluna contém principalmente números sequenciais
            try:
                valores_numericos = pd.to_numeric(df[col], errors='coerce').dropna()
                if len(valores_numericos) > len(df) * 0.7:  # 70% dos valores são numéricos
                    colunas_detectadas['numero'] = col
                    break
            except:
                continue
    
    # Se não encontrou descrição, pegar a coluna com mais texto
    if not colunas_detectadas['descricao']:
        max_texto_col = None
        max_texto_len = 0
        
        for col in df.columns:
            if col != colunas_detectadas['numero']:
                texto_total = df[col].astype(str).str.len().sum()
                if texto_total > max_texto_len:
                    max_texto_len = texto_total
                    max_texto_col = col
        
        colunas_detectadas['descricao'] = max_texto_col
    
    log_debug(f"Colunas detectadas: {colunas_detectadas}")
    return colunas_detectadas

def extrair_texto_pdf(file_path: str) -> Dict[str, Any]:
    """
    Usa pdfplumber para extração de texto e camelot para extração de tabelas de PDFs.
    
    Args:
        file_path: Caminho para o arquivo PDF
        
    Returns:
        Dicionário com texto extraído e tabelas
    """
    resultado = {'texto': '', 'tabelas': [], 'erro': None}
    
    try:
        log_info(f"Extraindo PDF: {os.path.basename(file_path)}", "📄")
        
        # Extração de texto com pdfplumber
        with pdfplumber.open(file_path) as pdf:
            texto_completo = ""
            for page in pdf.pages:
                texto_pagina = page.extract_text()
                if texto_pagina:
                    texto_completo += texto_pagina + "\n"
            
            resultado['texto'] = texto_completo
        
        # Extração de tabelas com camelot (se disponível)
        if CAMELOT_AVAILABLE:
            try:
                # Tentar com lattice primeiro
                tables = camelot.read_pdf(file_path, flavor='lattice', **CAMELOT_SETTINGS)
                
                # Se não encontrar tabelas, tentar com stream
                if len(tables) == 0:
                    tables = camelot.read_pdf(file_path, flavor='stream')
                
                for table in tables:
                    if not table.df.empty:
                        resultado['tabelas'].append(table.df)
                
                log_success(f"Extraídas {len(resultado['tabelas'])} tabelas do PDF")
                
            except Exception as e:
                log_warning(f"Erro na extração de tabelas com camelot: {str(e)}")
        
    except Exception as e:
        resultado['erro'] = str(e)
        log_error(f"Erro ao extrair PDF {file_path}: {str(e)}")
    
    return resultado

def extrair_texto_docx(file_path: str) -> Dict[str, Any]:
    """
    Usa python-docx para extrair texto e tabelas de arquivos DOCX.
    
    Args:
        file_path: Caminho para o arquivo DOCX
        
    Returns:
        Dicionário com texto extraído e tabelas
    """
    resultado = {'texto': '', 'tabelas': [], 'erro': None}
    
    if not DOCX_AVAILABLE:
        resultado['erro'] = "python-docx não está disponível"
        return resultado
    
    try:
        log_info(f"Extraindo DOCX: {os.path.basename(file_path)}", "📝")
        
        doc = Document(file_path)
        
        # Extrair texto dos parágrafos
        texto_completo = ""
        for paragraph in doc.paragraphs:
            texto_completo += paragraph.text + "\n"
        
        resultado['texto'] = texto_completo
        
        # Extrair tabelas
        for table in doc.tables:
            dados_tabela = []
            for row in table.rows:
                linha_dados = []
                for cell in row.cells:
                    linha_dados.append(cell.text.strip())
                dados_tabela.append(linha_dados)
            
            if dados_tabela:
                # Converter para DataFrame
                try:
                    df_tabela = pd.DataFrame(dados_tabela[1:], columns=dados_tabela[0])
                    resultado['tabelas'].append(df_tabela)
                except:
                    # Se falhar, criar DataFrame sem cabeçalho
                    df_tabela = pd.DataFrame(dados_tabela)
                    resultado['tabelas'].append(df_tabela)
        
        log_success(f"Extraídas {len(resultado['tabelas'])} tabelas do DOCX")
        
    except Exception as e:
        resultado['erro'] = str(e)
        log_error(f"Erro ao extrair DOCX {file_path}: {str(e)}")
    
    return resultado

def extrair_texto_doc(file_path: str) -> Dict[str, Any]:
    """
    Tenta lidar com arquivos DOC legados (pode requerer conversão ou pular com aviso).
    
    Args:
        file_path: Caminho para o arquivo DOC
        
    Returns:
        Dicionário com resultado da extração
    """
    resultado = {'texto': '', 'tabelas': [], 'erro': None}
    
    log_warning(f"Arquivo DOC detectado: {os.path.basename(file_path)}")
    log_warning("Arquivos DOC legados não são suportados diretamente.")
    log_info("Sugestão: Converta para DOCX manualmente para melhor compatibilidade.")
    
    resultado['erro'] = "Formato DOC não suportado - converta para DOCX"
    return resultado

def extrair_texto_excel(file_path: str) -> Dict[str, Any]:
    """
    Usa pandas para ler arquivos Excel e extrair planilhas relevantes.
    
    Args:
        file_path: Caminho para o arquivo Excel
        
    Returns:
        Dicionário com dados extraídos
    """
    resultado = {'texto': '', 'tabelas': [], 'erro': None}
    
    try:
        log_info(f"Extraindo Excel: {os.path.basename(file_path)}", "📊")
        
        # Ler todas as planilhas
        excel_file = pd.ExcelFile(file_path)
        
        for sheet_name in excel_file.sheet_names:
            try:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                if not df.empty:
                    resultado['tabelas'].append(df)
                    
                    # Adicionar nome da planilha como texto
                    resultado['texto'] += f"PLANILHA: {sheet_name}\n"
                    
            except Exception as e:
                log_warning(f"Erro ao ler planilha '{sheet_name}': {str(e)}")
        
        log_success(f"Extraídas {len(resultado['tabelas'])} planilhas do Excel")
        
    except Exception as e:
        resultado['erro'] = str(e)
        log_error(f"Erro ao extrair Excel {file_path}: {str(e)}")
    
    return resultado

def extrair_texto_txt(file_path: str) -> Dict[str, Any]:
    """
    Leitura simples de arquivos de texto com detecção de encoding.
    
    Args:
        file_path: Caminho para o arquivo TXT
        
    Returns:
        Dicionário com texto extraído
    """
    resultado = {'texto': '', 'tabelas': [], 'erro': None}
    
    try:
        log_info(f"Extraindo TXT: {os.path.basename(file_path)}", "📄")
        
        # Tentar detectar encoding
        encoding_usado = 'utf-8'
        
        if CHARDET_AVAILABLE:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding_detectado = chardet.detect(raw_data)
                if encoding_detectado['encoding']:
                    encoding_usado = encoding_detectado['encoding']
        
        # Tentar ler com diferentes encodings
        for encoding in [encoding_usado] + ENCODING_OPTIONS:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    resultado['texto'] = f.read()
                    log_success(f"Arquivo lido com encoding: {encoding}")
                    break
            except UnicodeDecodeError:
                continue
        
        if not resultado['texto']:
            resultado['erro'] = "Não foi possível decodificar o arquivo com nenhum encoding"
        
    except Exception as e:
        resultado['erro'] = str(e)
        log_error(f"Erro ao extrair TXT {file_path}: {str(e)}")
    
    return resultado

def processar_arquivo_zip(file_path: str) -> Dict[str, Any]:
    """
    Extrai conteúdo de ZIP para diretório temporário e processa arquivos internos recursivamente.
    
    Args:
        file_path: Caminho para o arquivo ZIP
        
    Returns:
        Dicionário com resultados da extração
    """
    resultado = {'texto': '', 'tabelas': [], 'erro': None, 'arquivos_processados': []}
    
    try:
        log_info(f"Processando ZIP: {os.path.basename(file_path)}", "📦")
        
        # Criar diretório temporário
        with tempfile.TemporaryDirectory() as temp_dir:
            # Extrair ZIP
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # Processar arquivos extraídos recursivamente
            arquivos_extraidos = percorrer_diretorios(temp_dir)
            
            for arquivo in arquivos_extraidos:
                caminho_arquivo = arquivo['caminho_completo']
                extensao = arquivo['extensao']
                
                # Processar cada arquivo baseado na extensão
                resultado_arquivo = extrair_descricoes_complementares(caminho_arquivo)
                
                if resultado_arquivo and not resultado_arquivo.get('erro'):
                    resultado['texto'] += f"\n--- ARQUIVO: {arquivo['nome_arquivo']} ---\n"
                    resultado['texto'] += resultado_arquivo.get('texto', '')
                    resultado['tabelas'].extend(resultado_arquivo.get('tabelas', []))
                    resultado['arquivos_processados'].append(arquivo['nome_arquivo'])
        
        log_success(f"ZIP processado: {len(resultado['arquivos_processados'])} arquivos internos")
        
    except Exception as e:
        resultado['erro'] = str(e)
        log_error(f"Erro ao processar ZIP {file_path}: {str(e)}")
    
    return resultado

def buscar_tabelas_complementares(texto: str, tabelas: List[pd.DataFrame]) -> List[pd.DataFrame]:
    """
    Busca por tabelas contendo palavras-chave relevantes usando padrões regex.
    
    Args:
        texto: Texto extraído do documento
        tabelas: Lista de DataFrames das tabelas extraídas
        
    Returns:
        Lista de tabelas relevantes
    """
    tabelas_relevantes = []
    
    # Buscar no texto por palavras-chave
    texto_lower = texto.lower()
    keywords_encontradas = []
    
    for keyword in TABLE_KEYWORDS:
        if re.search(keyword, texto_lower, re.IGNORECASE):
            keywords_encontradas.append(keyword)
    
    log_debug(f"Keywords encontradas no texto: {keywords_encontradas}")
    
    # Analisar tabelas
    for i, tabela in enumerate(tabelas):
        tabela_relevante = False
        
        # Verificar se a tabela contém colunas de item e descrição
        colunas_detectadas = detectar_colunas_tabela(tabela)
        
        if colunas_detectadas['numero'] and colunas_detectadas['descricao']:
            tabela_relevante = True
            log_debug(f"Tabela {i+1} identificada como relevante (contém item + descrição)")
        
        # Verificar conteúdo da tabela por keywords
        if not tabela_relevante:
            tabela_texto = ' '.join(tabela.astype(str).values.flatten()).lower()
            for keyword in TABLE_KEYWORDS:
                if re.search(keyword, tabela_texto, re.IGNORECASE):
                    tabela_relevante = True
                    log_debug(f"Tabela {i+1} identificada como relevante (keyword: {keyword})")
                    break
        
        if tabela_relevante:
            tabelas_relevantes.append(tabela)
    
    log_info(f"Encontradas {len(tabelas_relevantes)} tabelas relevantes de {len(tabelas)} total")
    return tabelas_relevantes

def extrair_descricoes_complementares(file_path: str) -> Dict[str, Any]:
    """
    Função principal de extração que despacha para o extrator apropriado baseado no tipo de arquivo.
    
    Args:
        file_path: Caminho para o arquivo
        
    Returns:
        Dicionário com dados extraídos
    """
    if not os.path.exists(file_path):
        return {'erro': f"Arquivo não encontrado: {file_path}"}
    
    extensao = Path(file_path).suffix.lower()
    
    # Despachar para função apropriada
    if extensao == '.pdf':
        return extrair_texto_pdf(file_path)
    elif extensao == '.docx':
        return extrair_texto_docx(file_path)
    elif extensao == '.doc':
        return extrair_texto_doc(file_path)
    elif extensao in ['.xlsx', '.xls']:
        return extrair_texto_excel(file_path)
    elif extensao == '.txt':
        return extrair_texto_txt(file_path)
    elif extensao == '.zip':
        return processar_arquivo_zip(file_path)
    else:
        return {'erro': f"Extensão não suportada: {extensao}"}

def fazer_matching_itens(df_master: pd.DataFrame, df_complementares: pd.DataFrame) -> pd.DataFrame:
    """
    Faz merge dos dados master com descrições complementares usando números de itens normalizados.
    
    Args:
        df_master: DataFrame com dados do master.xlsx
        df_complementares: DataFrame com descrições complementares extraídas
        
    Returns:
        DataFrame merged com coluna DESCRICAO_REFERENCIA
    """
    log_info("Iniciando processo de matching de itens", "🔗")
    
    # Criar cópia do master para não modificar o original
    df_resultado = df_master.copy()
    
    # Adicionar coluna DESCRICAO_REFERENCIA vazia
    df_resultado['DESCRICAO_REFERENCIA'] = ''
    
    if df_complementares.empty:
        log_warning("Nenhuma descrição complementar encontrada para matching")
        return df_resultado
    
    # Normalizar números de itens em ambos os DataFrames
    df_master_norm = df_master.copy()
    df_master_norm['NUMERO_NORMALIZADO'] = df_master_norm['Nº'].apply(normalizar_numero_item)
    
    df_comp_norm = df_complementares.copy()
    if 'numero_item' in df_comp_norm.columns:
        df_comp_norm['NUMERO_NORMALIZADO'] = df_comp_norm['numero_item'].apply(normalizar_numero_item)
    else:
        log_warning("Coluna 'numero_item' não encontrada nas descrições complementares")
        return df_resultado
    
    # Fazer matching
    matches_encontrados = 0
    
    for idx, row_master in df_master_norm.iterrows():
        numero_master = row_master['NUMERO_NORMALIZADO']
        arquivo_master = str(row_master.get('ARQUIVO', ''))
        
        if not numero_master:
            continue
        
        # Buscar correspondência nas descrições complementares
        matches = df_comp_norm[df_comp_norm['NUMERO_NORMALIZADO'] == numero_master]
        
        if not matches.empty:
            # Se encontrou match, usar a primeira descrição encontrada
            descricao_complementar = matches.iloc[0].get('descricao', '')
            
            if descricao_complementar:
                df_resultado.loc[idx, 'DESCRICAO_REFERENCIA'] = str(descricao_complementar)
                matches_encontrados += 1
                log_debug(f"Match encontrado - Item {numero_master}: {descricao_complementar[:50]}...")
    
    log_success(f"Matching concluído: {matches_encontrados} itens com descrições complementares")
    
    return df_resultado

def gerar_master_heavy(df_merged: pd.DataFrame, output_path: str) -> bool:
    """
    Cria o arquivo Excel final com todas as colunas originais mais DESCRICAO_REFERENCIA.
    
    Args:
        df_merged: DataFrame com dados merged
        output_path: Caminho para salvar o arquivo
        
    Returns:
        True se sucesso, False caso contrário
    """
    try:
        log_info("Gerando arquivo master_heavy.xlsx", "📊")
        
        # Criar diretório de saída se não existir
        output_dir = os.path.dirname(output_path)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            log_info(f"Diretório criado: {output_dir}")
        
        # Salvar arquivo Excel
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            df_merged.to_excel(writer, index=False, sheet_name='Master_Heavy')
            
            # Ajustar largura das colunas
            worksheet = writer.sheets['Master_Heavy']
            
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                
                adjusted_width = min(max_length + 2, 100)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        log_success(f"Arquivo salvo: {output_path}")
        log_info(f"Total de linhas: {len(df_merged)}")
        
        # Estatísticas
        itens_com_referencia = len(df_merged[df_merged['DESCRICAO_REFERENCIA'].str.len() > 0])
        log_info(f"Itens com descrição de referência: {itens_com_referencia}/{len(df_merged)}")
        
        return True
        
    except Exception as e:
        log_error(f"❌ Erro ao gerar arquivo final: {str(e)}")
        return False

# ======================================================================
# FUNÇÃO PRINCIPAL
# ======================================================================

def main():
    """Função principal que executa todo o pipeline de extração."""
    
    print("=" * 80)
    print("🎯 EXTRATOR DE EDITAIS COMPLETO")
    print("=" * 80)
    
    # Verificar se os diretórios e arquivos existem
    log_info("Verificando configuração inicial...")
    
    if not os.path.exists(EDITAIS_DIR):
        log_error(f"Diretório de editais não encontrado: {EDITAIS_DIR}")
        return False
    
    if not os.path.exists(MASTER_XLSX_PATH):
        log_error(f"Arquivo master.xlsx não encontrado: {MASTER_XLSX_PATH}")
        return False
    
    log_success("Configuração inicial verificada")
    
    # Carregar arquivo master
    try:
        log_info("Carregando arquivo master.xlsx...")
        df_master = pd.read_excel(MASTER_XLSX_PATH)
        log_success(f"Master carregado: {len(df_master)} itens")
        
        # Verificar colunas obrigatórias
        colunas_obrigatorias = ['ARQUIVO', 'Nº', 'DESCRICAO']
        colunas_faltantes = [col for col in colunas_obrigatorias if col not in df_master.columns]
        
        if colunas_faltantes:
            log_error(f"Colunas obrigatórias faltantes no master: {colunas_faltantes}")
            return False
            
    except Exception as e:
        log_error(f"❌ Erro ao carregar master.xlsx: {str(e)}")
        return False
    
    # Percorrer diretório de editais
    arquivos_editais = percorrer_diretorios(EDITAIS_DIR)
    
    if not arquivos_editais:
        log_warning("Nenhum arquivo de edital encontrado para processar")
        return False
    
    # Processar cada arquivo e extrair descrições complementares
    log_info("Iniciando extração de descrições complementares...")
    
    todas_descricoes = []
    arquivos_processados = 0
    arquivos_com_erro = 0
    
    for arquivo_info in arquivos_editais:
        caminho_arquivo = arquivo_info['caminho_completo']
        nome_arquivo = arquivo_info['nome_arquivo']
        
        try:
            log_info(f"Processando: {nome_arquivo}")
            
            # Extrair dados do arquivo
            resultado_extracao = extrair_descricoes_complementares(caminho_arquivo)
            
            if resultado_extracao.get('erro'):
                log_warning(f"Erro na extração: {resultado_extracao['erro']}")
                arquivos_com_erro += 1
                continue
            
            # Buscar tabelas relevantes
            texto = resultado_extracao.get('texto', '')
            tabelas = resultado_extracao.get('tabelas', [])
            
            tabelas_relevantes = buscar_tabelas_complementares(texto, tabelas)
            
            # Processar tabelas relevantes
            for tabela in tabelas_relevantes:
                colunas_detectadas = detectar_colunas_tabela(tabela)
                
                if colunas_detectadas['numero'] and colunas_detectadas['descricao']:
                    # Extrair dados da tabela
                    for _, row in tabela.iterrows():
                        numero_item = row[colunas_detectadas['numero']]
                        descricao = row[colunas_detectadas['descricao']]
                        
                        if pd.notna(numero_item) and pd.notna(descricao):
                            descricao_info = {
                                'arquivo_origem': nome_arquivo,
                                'numero_item': numero_item,
                                'descricao': str(descricao).strip()
                            }
                            todas_descricoes.append(descricao_info)
            
            arquivos_processados += 1
            
            # Pequena pausa para não sobrecarregar o sistema
            time.sleep(0.1)
            
        except Exception as e:
            log_error(f"Erro ao processar {nome_arquivo}: {str(e)}")
            arquivos_com_erro += 1
            continue
    
    log_success(f"Extração concluída: {arquivos_processados} arquivos processados")
    log_info(f"Arquivos com erro: {arquivos_com_erro}")
    log_info(f"Total de descrições extraídas: {len(todas_descricoes)}")
    
    # Converter descrições para DataFrame
    if todas_descricoes:
        df_complementares = pd.DataFrame(todas_descricoes)
        log_success("DataFrame de descrições complementares criado")
    else:
        log_warning("Nenhuma descrição complementar foi extraída")
        df_complementares = pd.DataFrame()
    
    # Fazer matching com o master
    df_final = fazer_matching_itens(df_master, df_complementares)
    
    # Gerar arquivo final
    sucesso = gerar_master_heavy(df_final, OUTPUT_PATH)
    
    if sucesso:
        print("\n" + "=" * 80)
        log_success("✅ PROCESSO CONCLUÍDO COM SUCESSO!")
        log_info(f"Arquivo gerado: {OUTPUT_PATH}")
        print("=" * 80)
        return True
    else:
        log_error("✅ PROCESSO FINALIZADO COM ERROS")
        return False

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        log_warning("Processo interrompido pelo usuário")
    except Exception as e:
        log_error(f"Erro inesperado: {str(e)}")
        traceback.print_exc()